from docx import Document
from datetime import datetime
import pandas as pd

# lendo o arquivo excel através do pandas
tabela = pd.read_excel('Informações.xlsx') # quando estiver fora da pasta do projeto digitar o caminho completo
# print(tabela)

for linha in tabela.index:
    # lendo o arquivo do word
    documento = Document('Contrato.docx')  # quando estiver fora da pasta do projeto digitar o caminho completo

    # print(documento.paragraphs)
    nome = tabela.loc[linha, 'Nome']
    item1 = tabela.loc[linha, 'Item1']
    item2 = tabela.loc[linha, 'Item2']
    item3 = tabela.loc[linha, 'Item 3']

    # criando um dicionário
    referencias = {
        'XXXX': nome,
        'YYYY': item1,
        'ZZZZ': item2,
        'WWWW': item3,
        'DD': str(datetime.now().day),
        'MM': str(datetime.now().month),
        'AAAA': str(datetime.now().year)
    }

    for paragrafo in documento.paragraphs:
        for codigo in referencias:
            valor = referencias[codigo]
            paragrafo.text = paragrafo.text.replace(codigo, valor)

    # Alterando o nome de quem assina
    documento.paragraphs[14].text = nome

    # for paragrafo in documento.paragraphs:
    #     print(paragrafo.text)
    #     print('-')

    documento.save(f'Contrato - {nome}.docx')